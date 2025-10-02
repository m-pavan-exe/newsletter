"""
Output module for exporting summary data to Word and PDF.
"""
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
import os
import subprocess
from typing import Dict, Any
from utils import get_logger

logger = get_logger(__name__)
import re

# Color constants
PRIMARY_COLOR_HEX = '#AE2278'
PRIMARY_COLOR = RGBColor(0xAE, 0x22, 0x78)

# ---------------------------------------------------------------------------
# Helper Utilities
# ---------------------------------------------------------------------------

def _safe_makedirs(path: str) -> None:
    """Safely create *path* directory, logging any errors."""
    try:
        os.makedirs(path, exist_ok=True)
    except OSError as exc:
        logger.error("Failed to create directory %s: %s", path, exc)

def replace_table_placeholders(doc: Document, summary_dict: Dict[str, Any]) -> None:
    placeholder_pattern = re.compile(r"{([a-zA-Z0-9_]+)}")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = placeholder_pattern.findall(cell.text)
                for key in matches:
                    if key in summary_dict:
                        placeholder = f"{{{key}}}"
                        cell.text = cell.text.replace(placeholder, str(summary_dict[key]))

def _add_section_header(doc, text):
    """Add a colored section header to the document."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = PRIMARY_COLOR
    return para

def _add_subsection_header(doc, text):
    """Add a subsection header with minimal spacing."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    return para

def _add_paragraph(doc, text, indent=0, bold=False, italic=False):
    """Add a formatted paragraph to the document."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return para

def _process_business_lines(doc, summary_dict, prefix, header_text):
    """Process and add business line sections (overcharge/undercharge)."""
    header_key = f"{prefix}_header"
    aggs_key = f"{prefix}_aggs"
    summaries_key = f"{prefix}_summary"
    recs_key = f"{prefix}_recommendations"
    
    if header_key in summary_dict:
        # Add section header
        _add_section_header(doc, header_text)
        
        # Add main header content if exists
        if summary_dict[header_key]:
            _add_paragraph(doc, summary_dict[header_key])
        
        # Process each item in the business line
        aggs = summary_dict.get(aggs_key, [])
        summaries = summary_dict.get(summaries_key, [])
        recommendations = summary_dict.get(recs_key, [])
        
        for idx, (agg, summary, rec) in enumerate(zip(aggs, summaries, recommendations), 1):
            # Add aggregation point
            agg = agg.rstrip("\n\n")
            _add_subsection_header(doc, f"{idx}. {agg}")
            
            # Add summary if exists
            if idx <= len(summaries) and summaries[idx-1]:
                summary_para = doc.add_paragraph()
                summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                summary_para.paragraph_format.space_before = Pt(0)  # Remove space before paragraph
                
                # Add bold label and text in the same paragraph
                run_label = summary_para.add_run("Summary: \n")
                run_label.bold = True
                run_label.font.size = Pt(12)
                summary_para.add_run(summaries[idx-1]).font.size = Pt(12)
            
            # Add recommendation if exists
            if idx <= len(recommendations) and recommendations[idx-1]:
                rec_label = doc.add_paragraph()
                rec_label.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                run_label = rec_label.add_run("Recommendation:\n")
                run_label.bold = True
                run_label.font.size = Pt(12)
                rec_label.add_run(recommendations[idx-1]).font.size = Pt(12)

def export_ra_ai_summary_to_pdf(summary_dict: Dict[str, Any], template_path: str, output_path: str) -> None:
    try:
        doc = Document(template_path)
        
        # Format and set the date
        run_date = datetime.now().strftime('%d-%B-%Y')
        day, month, year = run_date.split('-')
        formatted_date = f"{int(day)}-{month.capitalize()}-{year}"
        summary_dict['date'] = formatted_date
        
        # Replace placeholders in tables and date in runs
        replace_table_placeholders(doc, summary_dict)
        
        # Replace date placeholder in all runs
        for para in doc.paragraphs:
            for run in para.runs:
                if '{date}' in run.text:
                    run.text = run.text.replace('{date}', summary_dict.get('date', ''))

        
        # Add Summary section if it exists
        if 'segments_summary' in summary_dict and summary_dict['segments_summary']:
            # Add Summary header
            _add_section_header(doc, "Summary")
            
            # Add summary content
            _add_paragraph(doc, summary_dict['segments_summary'])
            
        
        # Process Overcharged SOC Lines
        _process_business_lines(doc, summary_dict, 'overcharge_business_line', 'Overcharged SOC Lines')
        
        # Process Undercharged SOC Lines
        _process_business_lines(doc, summary_dict, 'undercharge_business_line', 'Undercharged SOC Lines')
        


        doc.save(output_path)
        output_dir = os.path.dirname(output_path)
        _safe_makedirs(output_dir)
        subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", output_dir,
            output_path
        ], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        pdf_path = os.path.join(output_dir, os.path.splitext(os.path.basename(output_path))[0] + '.pdf')
        logger.info("PDF created at: %s", pdf_path)
        # Remove DOCX only if PDF conversion succeeded
        if os.path.exists(output_path):
            try:
                os.remove(output_path)
                logger.info("Removed DOCX file after PDF conversion: %s", output_path)
            except OSError as exc:
                logger.error("Failed to remove DOCX file: %s", exc)
    except subprocess.CalledProcessError as exc:
        logger.error("Failed to convert Word to PDF: %s", exc)
    except Exception as e:
        logger.error("Error in export_ra_ai_summary_to_pdf: %s", e)
