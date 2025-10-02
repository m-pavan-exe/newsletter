import os
from docx import Document
from datetime import datetime
from typing import Dict, List, Any
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from pathlib import Path
import re
import subprocess
from utils import get_logger

logger = get_logger(__name__)

PRIMARY_COLOR_HEX = '#AE2278'
PRIMARY_COLOR = RGBColor(0xAE, 0x22, 0x78)

def export_bs_summary_to_docx_v4(
    template_path: str,
    output_dir: str,
    summaries_dict: Dict[str, str],
    source_level_summaries: List[Dict[str, Any]]
) -> str:
    """Generate a Word document (v4 format) for the BS AI summary.
    Modified to include separate BS and IS tables and summaries.
    """
    doc = Document(template_path)

    # ----------------------------------------------------------
    # 1. Replace {date} placeholder
    # ----------------------------------------------------------
    formatted_date = datetime.today().strftime("%d-%B-%Y")
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if "{date}" in run.text:
                run.text = run.text.replace("{date}", formatted_date)

    # ----------------------------------------------------------
    # 2. Replace {summary} placeholder with executive sections
    # ----------------------------------------------------------
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if "{summary}" in run.text:
                # Remove the entire paragraph containing the placeholder
                p = paragraph._element
                p.getparent().remove(p)

                # ----------------------------------------------------------
                # 2.1 Executive Summary
                # ----------------------------------------------------------
                header_para = doc.add_paragraph("Executive Summary")
                header_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                header_para.paragraph_format.space_before = Pt(6)
                header_para.paragraph_format.space_after = Pt(6)
                header_run = header_para.runs[0]
                header_run.bold = True
                header_run.font.size = Pt(12)
                header_run.font.name = "Alptos"
                header_run.font.color.rgb = PRIMARY_COLOR

                # Introduction summary
                summary_text = summaries_dict.get("INTRODUCTION", "").strip("\n")
                if summary_text:
                    summary_para = doc.add_paragraph("")
                    summary_run = summary_para.add_run(summary_text)
                    summary_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    summary_run.font.size = Pt(10)
                    summary_run.font.name = "Alptos"

                # ----------------------------------------------------------
                # 2.2 Balance Sheet Analysis
                # ----------------------------------------------------------
                bs_items = [tl for tl in source_level_summaries 
                          if tl.get("MGT_LINE_DESCRIPTION") in ["Total Assets", "Total Liability"]]
                
                if bs_items:
                    # Add Balance Sheet header
                    header_para = doc.add_paragraph("Balance Sheet Analysis")
                    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    header_para.paragraph_format.space_before = Pt(16)
                    header_para.paragraph_format.space_after = Pt(6)
                    header_run = header_para.runs[0]
                    header_run.bold = True
                    header_run.font.size = Pt(10)
                    header_run.font.name = "Alptos"
                    header_run.font.color.rgb = PRIMARY_COLOR

                    # Add Balance Sheet table
                    _add_summary_table(doc, bs_items, level='MGT_LINE_DESCRIPTION')
                    
                    # Add Balance Sheet summary text
                    bs_summary_text = summaries_dict.get("BS_SUMMARY", "").strip("\n")
                    if bs_summary_text:
                        _add_summary_paragraph(doc, bs_summary_text)
                    
                    # Add Balance Sheet details
                    _add_source_details_bs(doc, bs_items)

                # ----------------------------------------------------------
                # 2.3 Income Statement Analysis
                # ----------------------------------------------------------
                is_items = [tl for tl in source_level_summaries 
                          if tl.get("MGT_LINE_DESCRIPTION") in ["Total Income", "Total Expense"]]
                
                if is_items:
                    # Add Income Statement header
                    header_para = doc.add_paragraph("Income Statement Analysis")
                    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    header_para.paragraph_format.space_before = Pt(16)
                    header_para.paragraph_format.space_after = Pt(6)
                    header_run = header_para.runs[0]
                    header_run.bold = True
                    header_run.font.size = Pt(10)
                    header_run.font.name = "Alptos"
                    header_run.font.color.rgb = PRIMARY_COLOR

                    # Add Income Statement table with MTD
                    _add_summary_table(doc, is_items, level='MGT_LINE_DESCRIPTION', 
                                     mtd_level='TL_MTD_FIX', mtd_flag=True)
                    
                    # Add Income Statement summary text
                    is_summary_text = summaries_dict.get("IS_SUMMARY", "").strip("\n")
                    if is_summary_text:
                        _add_summary_paragraph(doc, is_summary_text)
                    
                    # Add Income Statement details
                    _add_source_details_is(doc, is_items)

   
    # ------------------------------------------------------------------
    # 4. Save
    # ------------------------------------------------------------------
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    out_path = Path(output_dir) / f"bs_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(str(out_path))
    return str(out_path)


def _add_summary_table(doc, items, level='MGT_LINE_DESCRIPTION', mtd_level='TL_MTD_FIX', mtd_flag=False):
    """Helper to add a neatly formatted summary table with full borders and column headers."""
    num_cols = 4 if mtd_flag else 3
    table = doc.add_table(rows=0, cols=num_cols)
    table.style = "Table Grid"
    table.autofit = True
    tbl = table._element

    # Set all table borders
    tbl_borders = parse_xml(r'''
    <w:tblBorders %s>
        <w:top w:val="single" w:sz="4" w:color="DDDDDD"/>
        <w:left w:val="single" w:sz="4" w:color="DDDDDD"/>
        <w:bottom w:val="single" w:sz="4" w:color="DDDDDD"/>
        <w:right w:val="single" w:sz="4" w:color="DDDDDD"/>
        <w:insideH w:val="single" w:sz="4" w:color="DDDDDD"/>
        <w:insideV w:val="single" w:sz="4" w:color="DDDDDD"/>
    </w:tblBorders>
    ''' % nsdecls('w'))
    tbl.tblPr.append(tbl_borders)

    # Add header row
    header_cells = table.add_row().cells
    header_titles = [""]
    if mtd_flag:
        header_titles += ["MTD", "Today", "Previous"]
    else:
        header_titles += ["Today", "Previous"]

    for idx, title in enumerate(header_titles):
        cell = header_cells[idx]
        cell.text = title
        p = cell.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = p.runs[0]
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Add data rows
    for tl in items:
        desc = str(tl.get(level, ""))
        today = str(tl.get("BUSINESS_DAY_FIX", ""))
        prev = str(tl.get("PREVIOUS_DAY_FIX", ""))
        if mtd_flag:
            mtd = str(tl.get(mtd_level, ""))
            aggs_list = [mtd, today, prev]
        else:
            aggs_list = [today, prev]

        row_cells = table.add_row().cells
        row_cells[0].text = desc
        p = row_cells[0].paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xAE, 0x22, 0x78)

        for j, val in enumerate(aggs_list):
            col_idx = j + 1
            cell = row_cells[col_idx]
            cell.text = val
            p = cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = "Alptos"




def _add_summary_paragraph(doc, text):
    """Helper to add a summary paragraph."""
    para = doc.add_paragraph("")
    run = para.add_run(text)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run.font.size = Pt(10)
    run.font.name = "Alptos"


def _add_source_details_is(doc, items):
    """Helper to add detailed analysis sections for items (numbered sentences)."""
    for mgt_line in items:
        desc = str(mgt_line.get("MGT_LINE_DESCRIPTION", ""))
        header_para = doc.add_paragraph(f"{desc} Analysis Summary")
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        header_run = header_para.runs[0]
        header_run.bold = True
        header_run.font.size = Pt(10)
        header_run.font.name = "Alptos"
        header_run.font.color.rgb = PRIMARY_COLOR
        header_para.paragraph_format.space_before = Pt(11)
        header_para.paragraph_format.space_after = Pt(2)
        header_para.paragraph_format.keep_with_next = True

        significant_sources = mgt_line.get("SIGNIFICANT_SOURCES", [])
        if significant_sources:
            _add_summary_table(doc, significant_sources, level='SOURCE_MGT_LINE_DESC', mtd_level='SRC_MTD', mtd_flag=True)
            for src in significant_sources:
                summary_text = src.get("SUMMARY", "").lstrip("\n")
                if summary_text:
                    para = doc.add_paragraph()
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    # para.paragraph_format.left_indent = Pt(18)
                    para.paragraph_format.space_before = Pt(5)
                    para.paragraph_format.space_after = Pt(5)
                        
                    # Add index prefix
                    para.add_run("• ").bold = True  # Add bullet manually

                    # Process bold segments within the sentence
                    parts = re.split(r'(\*\*.*?\*\*)', summary_text)
                    for part in parts:
                        if not part:
                            continue
                        clean_part = part.replace('**', '')
                        run = para.add_run(clean_part)
                        run.font.size = Pt(10)
                        run.font.name = "Alptos"
                        if part.startswith('**') and part.endswith('**'):
                            run.bold = True
                    
                    # Add period if missing 
                    if not para.text.strip().endswith('.'):
                        para.add_run(".")

def _add_source_details_bs(doc, items):
    """Helper to add detailed analysis sections for items (numbered sentences)."""
    for mgt_line in items:
        desc = str(mgt_line.get("MGT_LINE_DESCRIPTION", ""))
        header_para = doc.add_paragraph(f"{desc} Analysis Summary")
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        header_run = header_para.runs[0]
        header_run.bold = True
        header_run.font.size = Pt(10)
        header_run.font.name = "Alptos"
        header_run.font.color.rgb = PRIMARY_COLOR
        header_para.paragraph_format.space_before = Pt(11)
        header_para.paragraph_format.space_after = Pt(2)
        header_para.paragraph_format.keep_with_next = True

        significant_sources = mgt_line.get("SIGNIFICANT_SOURCES", [])
        if significant_sources:
            _add_summary_table(doc, significant_sources, level='SOURCE_MGT_LINE_DESC')
            for src in significant_sources:
                summary_text = src.get("SUMMARY", "").lstrip("\n")
                if summary_text:
                    para = doc.add_paragraph()
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    # para.paragraph_format.left_indent = Pt(18)
                    para.paragraph_format.space_before = Pt(5)
                    para.paragraph_format.space_after = Pt(5)

                    para.add_run("• ").bold = True  # Add bullet manually
                    parts = re.split(r'(\*\*.*?\*\*)', summary_text)
                    for part in parts:
                        if not part:
                            continue
                        clean_part = part.replace('**', '')
                        run = para.add_run(clean_part)
                        run.font.size = Pt(10)
                        run.font.name = "Alptos"
                        if part.startswith('**') and part.endswith('**'):
                            run.bold = True

                    if not para.text.strip().endswith('.'):
                        para.add_run(".")
                        

def convert_docx_to_pdf(docx_path: str) -> str:
    """
    Convert a DOCX file to PDF using LibreOffice.
    
    Args:
        docx_path: Path to the input DOCX file
        
    Returns:
        str: Path to the converted PDF file if conversion was successful, empty string otherwise
    """
    try:
        output_dir = os.path.dirname(docx_path)
        
        # Create output directory if it doesn't exist
        _safe_makedirs(output_dir)
        
        # Convert the document using LibreOffice
        result = subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", output_dir,
            docx_path
        ], check=True, capture_output=True, text=True)
        
        # Check if conversion was successful
        if result.returncode == 0:
            pdf_path = os.path.join(output_dir, os.path.splitext(os.path.basename(docx_path))[0] + '.pdf')
            logger.info(f"PDF created at: {pdf_path}")
            
            # Remove the original DOCX file
            try:
                os.remove(docx_path)
                logger.info(f"Removed DOCX file after PDF conversion: {docx_path}")
            except Exception as e:
                logger.error(f"Failed to remove DOCX file: {e}")
                
            return pdf_path
        return False
        
    except subprocess.CalledProcessError as e:
        logger.error(f"LibreOffice conversion failed: {e.stderr}")
        return False
    except Exception as e:
        logger.error(f"Error converting DOCX to PDF: {str(e)}")
        return False

def _safe_makedirs(path: str) -> None:
    """Create directory *path* if it doesn't exist, logging errors."""
    try:
        os.makedirs(path, exist_ok=True)
    except OSError as exc:
        logger.error("Failed to create directory %s: %s", path, exc)



def export_bs_summary_to_html_v1(
    summaries_dict: Dict[str, str],
    source_level_summaries: List[Dict[str, Any]]
) -> str:
    """Generate an HTML string (v1 format) for the BS AI summary."""
    formatted_date = datetime.today().strftime("%d-%B-%Y")
    
    # Add date
    html = f"<strong>Date:</strong> {formatted_date}\n\n"
    
    # Add executive summary
    html += f"<strong>Executive Summary</strong>\n"
    intro_text = summaries_dict.get("INTRODUCTION", "").strip("\n")
    if intro_text:
        html += intro_text + "\n"
    
    # Add balance sheet analysis
    bs_items = [tl for tl in source_level_summaries if tl.get("MGT_LINE_DESCRIPTION") in ["Total Assets", "Total Liability"]]
    if bs_items:
        html += f"\n<strong>Balance Sheet Analysis</strong>\n"
        html += _generate_html_table(bs_items, mtd_flag=False)
        bs_summary = summaries_dict.get("BS_SUMMARY", "").strip("\n")
        if bs_summary:
            html += f"<strong>Summary</strong>:\n{bs_summary}\n\n"
        html += _generate_source_details_html(bs_items, mtd_flag=False)
    
    # Add income statement analysis
    is_items = [tl for tl in source_level_summaries if tl.get("MGT_LINE_DESCRIPTION") in ["Total Income", "Total Expense"]]
    if is_items:
        html += f"\n<strong>Income Statement Analysis</strong>\n"
        html += _generate_html_table(is_items, mtd_flag=True)
        is_summary = summaries_dict.get("IS_SUMMARY", "").strip("\n")
        if is_summary:
            html += f"<strong>Summary</strong>:\n{is_summary}\n\n"
        html += _generate_source_details_html(is_items, mtd_flag=True)
    
    return html


def _generate_html_table(items, mtd_flag=False):
    html = []
    html.append('<table class="tableAi">')
    
    headers = ['']
    headers += ['MTD', 'Today', 'Previous'] if mtd_flag else ['Today', 'Previous']
    
    html.append('<tr>' + ''.join(f'<th>{h}</th>' for h in headers) + '</tr>')
    
    for item in items:
        desc = item.get("MGT_LINE_DESCRIPTION", "") or item.get("SOURCE_MGT_LINE_DESC", "")
        today = item.get("BUSINESS_DAY_FIX", "")
        prev = item.get("PREVIOUS_DAY_FIX", "")
        mtd = item.get("TL_MTD_FIX", "") or item.get("SRC_MTD_FIX", "") if mtd_flag else None
        values = [mtd, today, prev] if mtd_flag else [today, prev]

        html.append('<tr>')
        html.append(f'<td>{desc}</td>')
        for val in values:
            html.append(f'<td>{val}</td>')
        html.append('</tr>')

    html.append('</table>')
    return "\n".join(html)


def _generate_source_details_html(items, mtd_flag=False):
    html = []
    for mgt_line in items:
        desc = mgt_line.get("MGT_LINE_DESCRIPTION", "")
        if desc:
            html.append(f"<strong>{desc} Analysis Summary</strong>")
        significant_sources = mgt_line.get("SIGNIFICANT_SOURCES", [])
        if significant_sources:
            html.append(_generate_html_table(significant_sources, mtd_flag=mtd_flag))
            for src in significant_sources:
                summary = src.get("SUMMARY", "").strip("\n")
                if summary:
                    # Convert **bold** to <strong>
                    parts = re.split(r'(\*\*.*?\*\*)', summary)
                    sentence = ""
                    for part in parts:
                        if not part:
                            continue
                        if part.startswith("**") and part.endswith("**"):
                            sentence += f'<strong>{part[2:-2]}</strong>'
                        else:
                            sentence += part
                    if not sentence.strip().endswith('.'):
                        sentence += '.'
                    html.append(f"• {sentence}\n")
    return "\n".join(html)